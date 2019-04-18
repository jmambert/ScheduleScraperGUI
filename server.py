import os
import json
from flask import Flask, render_template, request, jsonify, redirect, url_for
from flask_bootstrap import Bootstrap
from flask_login import current_user, login_user, LoginManager, UserMixin, logout_user, login_required
from scraper import infoFunction, grabData
import pandas
from openpyxl import load_workbook
from docx.text.paragraph import Paragraph
from docx import Document
app = Flask(__name__)
app.secret_key = os.urandom(24).hex()

login_manager = LoginManager()
login_manager.session_protection = "strong"
login_manager.init_app(app)
login_manager.login_view= "login"

J_USERNAME = 'jcampbell'
J_PASSWORD = 'sudoUser'

class User(UserMixin):
    username = None
    authenticated = False

    def __init__(self, username):
        self.username = username
        self.authenticated = True

    def is_active(self):
         return True

    def get_id(self):
        return self.username

    def is_authenticated(self):
        return self.authenticated

    def is_anonymous(self):
        return False

@login_manager.user_loader
def loader_user(username):
    user = User(username)
    return user

@app.route("/login", methods=["GET","POST"])
def login():
    if request.method == "POST" and "Username" in request.form:
        print("PrintStatement\n")
        if request.form['Username'] == J_USERNAME and request.form['Password'] == J_PASSWORD:
            logged_user = User(request.form['Username'])
            login_user(logged_user, remember=False)
            print("PrintStatement0\n")
            return redirect(url_for("index"))
        else:
            print("I'm different!")
            alertMessage = "Login Failed"
        #except Exception as e:
        print("a little bit more different than the other ones")
        alertMessage = "Login Failed"
        return render_template('login.html',login=True,alertmessage = alertMessage)
    return render_template('login.html',login=True) 

@app.route('/')
@login_required
def index():
    return render_template('index.html')

@app.route('/professorAddOfficeHours')
@login_required
def professorPage():
    return render_template('professorPage.html')

@app.route('/semesterSchedule')
@login_required
def semesterSchedule():
    data = infoFunction()
    print(data)
    data = str(data.to_html())
    print(type(data))
    print(data)
    professorVariable = "Professor"
    return render_template('semesterSchedule.html', professorName = data)

@app.route('/submit', methods=['POST'])
@login_required
def submit():
    scrapedData = grabData()
    data = json.loads(request.form['data'])
    info = data['info']
    print(info)
    document = Document('professorTemplate.docx')
    document.save(info['firstname'] + info['lastname'] + '.docx') 
    
    firstname = info['firstname']
    lastname = info['lastname']
    tables = document.tables

    officeHours = tables[0]
    classSchedule = tables[1]

    document.paragraphs[0].text = "%s %s" % (info['season'], info['year'])
    document.paragraphs[1].text = "%s %s" % (info['firstname'], info['lastname'])
    document.paragraphs[2].text = info['officeLocation']
    document.paragraphs[3].text = info['phoneNumber']
    document.paragraphs[4].text = info['email']

    info.pop('season')
    info.pop('year')
    info.pop('firstname')
    info.pop('lastname')
    info.pop('officeLocation')
    info.pop('phoneNumber')
    info.pop('email')
    timeList = {}
    dayList = {}
    print(info)
    for key,val in info.items():
        print("key is " + key)
        if "time" in key:
            timeList[key[4:]] = val
        else:
            dayList[key[3:]] = val
 
        mondayCell = tables[0].cell(2,0)
        tuesdayCell = tables[0].cell(2,1)
        wednesdayCell = tables[0].cell(2,2)
        thursdayCell = tables[0].cell(2,3)
        fridayCell = tables[0].cell(2,4)

    for key,val in dayList.items(): 
        if dayList[key] == "Monday":
            mondayCell.text = mondayCell.text + timeList[key] + "\n"
        elif dayList[key] == "Tuesday":
            tuesdayCell.text = tuesdayCell.text + timeList[key] + "\n"
        elif dayList[key] == "Wednesday":
            wednesdayCell.text = wednesdayCell.text + timeList[key] + "\n"
        elif dayList[key] == "Thursday":
            thursdayCell.text = thursdayCell.text + timeList[key] + "\n"
        elif dayList[key] == "Friday":
            fridayCell.text = fridayCell.text + timeList[key] + "\n"

    cellNum = 2

    for row in scrapedData:
        if lastname in row[0]:
            if cellNum != 2:
                tables[1].add_row()
            tables[1].cell(cellNum,0).text = tables[1].cell(cellNum,0).text + row[1]
            tables[1].cell(cellNum,1).text = tables[1].cell(cellNum,1).text + row[2]
            tables[1].cell(cellNum,2).text = tables[1].cell(cellNum,2).text + row[3]
            cellNum = cellNum + 1

    print(timeList)
    print(dayList)

    document.save(firstname + lastname + '.docx')

    return "Success"

if __name__ == '__main__':
    app.run()
    
